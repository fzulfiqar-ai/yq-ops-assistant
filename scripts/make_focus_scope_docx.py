"""Generate the Focus ERP data-scope Word document (docs/Focus ERP Data Scope - YQ Bahrain.docx).

Mirrors docs/focus-erp-data-scope.html. Regenerate after editing that page so the two agree:
    python scripts/make_focus_scope_docx.py
"""
from __future__ import annotations

import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor

ACCENT = RGBColor(0x0F, 0x64, 0x66)
INK = RGBColor(0x16, 0x23, 0x2B)
MUTED = RGBColor(0x5D, 0x6B, 0x70)
FLAG = RGBColor(0x9D, 0x44, 0x20)

DASH = "—"   # em dash
NDASH = "–"  # en dash
MID = "·"    # middle dot

d = Document()

for s in d.sections:
    s.left_margin = s.right_margin = Inches(0.9)
    s.top_margin = s.bottom_margin = Inches(0.8)

base = d.styles["Normal"]
base.font.name = "Calibri"
base.font.size = Pt(10.5)
base.paragraph_format.space_after = Pt(7)


def para(text="", size=10.5, bold=False, color=None, italic=False, after=7, before=0):
    p = d.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color is not None:
        r.font.color.rgb = color
    return p


def heading(text, size=15, before=16):
    return para(text, size=size, bold=True, color=INK, after=4, before=before)


def kv(p, label, value):
    a = p.add_run(label + "  ")
    a.bold = True
    a.font.size = Pt(9)
    b = p.add_run(value + "    ")
    b.font.size = Pt(9)
    b.font.color.rgb = MUTED


def table(headers, rows, widths, header_pt=8.5, body_pt=9, mono_col=None, bold_first=True):
    t = d.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(header_pt)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(val)
            r.font.size = Pt(body_pt)
            if i == 0 and bold_first:
                r.bold = True
            if mono_col is not None and i == mono_col:
                r.font.name = "Consolas"
                r.font.size = Pt(body_pt - 0.6)
    for w, col in zip(widths, t.columns):
        for cell in col.cells:
            cell.width = Inches(w)
    return t


# ---- masthead -------------------------------------------------------------
para("DATA ACCESS REQUEST  " + MID + "  YQ BAHRAIN MOBILE ACCESSORIES W.L.L",
     size=8.5, bold=True, color=ACCENT, after=4)
para("Focus ERP Data Scope", size=26, bold=True, color=INK, after=4)
para("What we need out of Focus, in what shape, and how often " + DASH +
     " so the right access can be quoted without guesswork.", size=12, color=MUTED, after=10)

p = d.add_paragraph()
p.paragraph_format.space_after = Pt(2)
kv(p, "To:", "Focus Softnet " + DASH + " account & technical team")
kv(p, "From:", "YQ Bahrain " + MID + " Furqan Ahmed")
p = d.add_paragraph()
p.paragraph_format.space_after = Pt(10)
kv(p, "Datasets:", "10")
kv(p, "Direction:", "Read-only (no write-back to Focus)")

# ---- 1. the ask -----------------------------------------------------------
heading("1.  What we are asking for", before=10)
para("We run an internal reporting portal on top of Focus data. Today a staff member logs into "
     "Focus every morning, runs eight reports by hand, downloads them, and uploads the files to "
     "our system. We want that to happen automatically.")
para("We are not asking to write anything back into Focus. Every dataset below is read-only. No "
     "document is created, edited, posted or cancelled by us " + DASH + " Focus stays the system "
     "of record and the only place transactions are entered.")
para("One thing to understand before quoting: our system already reads these exact reports and "
     "their exact columns. Whatever route we choose only has to deliver the same fields.")

# ---- 2. routes ------------------------------------------------------------
heading("2.  Three routes, in the order we would prefer them")
para("Ranked by how quickly they deliver value against how much they cost to build. We are open "
     "to any of them " + DASH + " please quote whichever you can actually support.", color=MUTED)

para("Route 1 " + DASH + " Scheduled report delivery  (preferred first step)",
     size=11.5, bold=True, color=ACCENT, after=3, before=8)
para("Focus runs the same eight reports on a schedule and drops the files somewhere we can "
     "collect them " + DASH + " SFTP, a watched folder, or an email inbox. Nothing about the "
     "report definitions changes.")
para("We prefer this first because it needs no integration work on our side at all: our system "
     "already parses these files. It is also the easiest for your team to support, because "
     "nothing new is exposed. Format: XLS/XLSX or CSV, unchanged.")

para("Route 2 " + DASH + " Read-only database access",
     size=11.5, bold=True, color=ACCENT, after=3, before=8)
para("A read-only login against the Focus database, limited to the sales, stock, receivables and "
     "pricing tables behind these reports. SELECT only; we would never write.")
para("This is the best long-term option for reporting: full history, no dependency on report "
     "layouts staying identical. Please tell us if this conflicts with your support terms " +
     DASH + " if granting it voids support, say so and we will drop this route.")

para("Route 3 " + DASH + " REST API", size=11.5, bold=True, color=ACCENT, after=3, before=8)
para("The Focus REST API, scoped to the datasets below, GET only. Most future-proof, and our "
     "choice if it covers the fields in section 3 and supports fetching only what changed.")
para("Our concern is coverage rather than capability: several of these are computed reporting "
     "figures " + DASH + " ageing buckets, profitability, price books " + DASH + " not plain "
     "transaction records. Please confirm the API returns these as computed values, or tell us "
     "which ones it cannot.")

# ---- 3. datasets ----------------------------------------------------------
d.add_page_break()
heading("3.  The data we need", before=0)
para("Ten datasets. The first eight are the reports we export by hand today; the last two are "
     "documents we key in manually. 'Matched on' is how we identify a record we already hold, so "
     "re-sending the same rows is harmless.", color=MUTED)

datasets = [
    ("Sales Day Book", "Revenue, per-item sales, daily chart",
     "invoice_no, line_no, date, customer_account, item_name, quantity, rate, gross, discount, "
     "taxable, vat_amount, total_amount, warehouse_name, narration",
     "Daily", "invoice_no + line_no"),
    ("Summary Sales Register", "Salesman performance, payment mode",
     "invoice_no, order_date, customer_name, gross, salesman, payment_mode, sales_account_name",
     "Daily", "invoice_no"),
    ("Stock Balance by Warehouse", "Stock on hand, value, low-stock alerts",
     "item_name, warehouse_name, net_qty, selling_rate, total_value, as_of_date",
     "Daily", "item + warehouse + date"),
    ("Stock Ledger", "Movements, branch transfers, reconciliation",
     "item_name, move_date, voucher, voucher_type, received_qty, received_rate, issued_qty, "
     "issued_rate, balance_qty, received_value, issued_value, balance_value, avg_rate, "
     "warehouse_name, to_warehouse_name, narration",
     "Daily", "voucher + item"),
    ("Customer Summary Ageing by Due Date", "Receivables, overdue chasing",
     "account, account_code, group_name, balance, ageing buckets (0" + NDASH + "30, 31" + NDASH +
     "60, 61" + NDASH + "90, 91" + NDASH + "120, 121" + NDASH + "150, 151" + NDASH + "180, 181" +
     NDASH + "210, over 210), total, last_receipt_date, as_of_date",
     "Daily", "account + as_of_date"),
    ("Product Profitability Report", "Margins, below-cost detection",
     "item_name, report_date, gross, discount_pct, net_amount, cogs, gross_profit, gp_margin_pct, "
     "misc_charges, net_profit, np_margin_pct",
     "Daily", "item + report_date"),
    ("MA Selling Price Book", "Standard (B2B) selling price",
     "item_name, sku_code, customer, warehouse, currency, start_date, end_date, min_qty, max_qty, price",
     "Weekly", "item + price book"),
    ("Modern Trade Seller Book", "Retail (B2C) selling price",
     "same fields as above", "Weekly", "item + price book"),
    ("Purchase Orders", "Order tracking, cost vs last order",
     "po_no, po_date, vendor, warehouse, line_no, item_code, description, qty, rate, gross",
     "On change", "po_no + line_no"),
    ("Goods Receipt / MRN", "Landed cost, receiving against PO",
     "mrn_no, mrn_date, po_no, vendor, item_code, description, qty, landed rate/cost",
     "On change", "mrn_no + line"),
]
table(["Dataset (as named in Focus)", "Used for", "Fields required", "Frequency", "Matched on"],
      datasets, [1.35, 1.15, 2.7, 0.62, 1.05], body_pt=7.9, mono_col=2)

para("All amounts in BHD. Dates in any unambiguous format, ideally YYYY-MM-DD. If a field above "
     "does not exist under that name in Focus, please tell us the equivalent rather than omitting "
     "it " + DASH + " every one of these is in use today.", size=9.5, before=8)

# ---- 4. gotchas -----------------------------------------------------------
heading("4.  Three things that will bite us if unaddressed")
para("These come from working with the current exports. Please confirm how each behaves in "
     "whichever route you quote.", color=MUTED)

gotchas = [
    ("1 " + MID + " Salesman is arriving in the wrong field",
     "In the Sales Day Book as currently configured, the salesman name comes through in the "
     "Warehouse Name column, not a salesman field. We have worked around it, but if the API or "
     "database exposes a proper salesman field we would rather use that. Please confirm which "
     "field genuinely holds the salesman."),
    ("2 " + MID + " Line totals are frequently empty",
     "total_amount on sales lines is often blank, so revenue has to be reconstructed from the "
     "other columns. Please confirm whether the API or database returns a reliably populated "
     "line total, and state clearly whether each amount is VAT-inclusive or VAT-exclusive. "
     "Getting this wrong misstates revenue."),
    ("3 " + MID + " We need to fetch only what changed",
     "Pulling full history every day is wasteful and slow. For Route 2 or 3 we need to request "
     "records by date range or a modified-since timestamp. If the API can only return everything, "
     "please say so up front " + DASH + " it changes which route makes sense."),
]
for title, body in gotchas:
    para(title, size=10.5, bold=True, color=FLAG, after=2, before=7)
    para(body, size=10)

# ---- 5. questions ---------------------------------------------------------
d.add_page_break()
heading("5.  Questions we need answered to decide", before=0)
para("Please answer against whichever route you recommend. Commercial answers matter as much as "
     "technical ones.", color=MUTED)

questions = [
    ("Which routes can you actually support?",
     "Of the three above, which are available on our licence and version " + DASH +
     " and which do you recommend?"),
    ("What is the cost, and what shape is it?",
     "One-time setup, annual licence, per-module, or per-call. Include anything that changes at renewal."),
    ("Is it read-only, and can you guarantee that?",
     "We want write access explicitly excluded, not merely unused."),
    ("How do we authenticate?",
     "API key, OAuth, or database credentials " + DASH + " and how are they rotated if compromised?"),
    ("Can we filter by date or changed-since?",
     "See point 3 above. This is close to a deal-breaker for the API route."),
    ("Are there rate limits or page sizes?", "And what happens when we hit them."),
    ("How far back can we pull?", "We want a one-time backfill of history, not just from go-live."),
    ("Is there a test environment?", "We will not test integrations against live accounting data."),
    ("Does the API return computed reporting figures?",
     "Specifically ageing buckets, profitability and price books " + DASH +
     " or only raw transactions we would have to recompute?"),
    ("What happens on a Focus upgrade?",
     "Who is responsible if a version change breaks the integration, and what notice do we get?"),
    ("How long to deliver?", "From order to working access, per route."),
]
for i, (q, sub) in enumerate(questions, 1):
    p = d.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(5)
    n = p.add_run("%02d   " % i)
    n.bold = True
    n.font.color.rgb = ACCENT
    n.font.size = Pt(9.5)
    qq = p.add_run(q)
    qq.bold = True
    qq.font.size = Pt(10.5)
    ps = d.add_paragraph()
    ps.paragraph_format.left_indent = Inches(0.42)
    ps.paragraph_format.space_after = Pt(3)
    rs = ps.add_run(sub)
    rs.font.size = Pt(9.5)
    rs.font.color.rgb = MUTED

# ---- 6. acceptance --------------------------------------------------------
heading("6.  What done looks like")
para("We consider this delivered when, for a full week and without anyone logging into Focus to "
     "run a report by hand:")
table(["Check", "Acceptance"], [
    ("All ten datasets arrive", "On their stated frequency, unattended"),
    ("Figures reconcile", "Revenue, stock value and receivables match the same reports run "
                          "manually in Focus, to the fillis"),
    ("Re-sends are safe", "Delivering the same rows twice does not duplicate or double-count"),
    ("Failures are visible", "A missed or failed delivery is detectable, not silent"),
    ("History is loaded", "Backfill completed to the agreed start date"),
], [1.7, 5.0])

para("Prepared by YQ Bahrain Mobile Accessories W.L.L for discussion with Focus Softnet. "
     "Read-only access only " + DASH + " no write-back to Focus is requested or required.",
     size=8.5, color=MUTED, before=14, italic=True)

OUT = os.path.join("docs", "Focus ERP Data Scope - YQ Bahrain.docx")
d.save(OUT)
print("saved: %s  (%.0f KB)" % (OUT, os.path.getsize(OUT) / 1024))
print("tables: %d" % len(d.tables))
